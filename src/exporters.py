from __future__ import annotations

from io import BytesIO

import pandas as pd


def dataframe_to_csv_bytes(frame: pd.DataFrame) -> bytes:
    return frame.to_csv(index=False).encode("utf-8-sig")


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("当前环境未安装 python-docx，无法导出 DOCX 报告。请执行：pip install python-docx") from exc

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(_strip_markdown_emphasis(line))

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def markdown_to_pdf_bytes(markdown_text: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    except ImportError as exc:
        raise RuntimeError("当前环境未安装 reportlab，无法导出 PDF 报告。请执行：pip install reportlab") from exc

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    output = BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=42,
        leftMargin=42,
        topMargin=42,
        bottomMargin=42,
    )
    sample = getSampleStyleSheet()
    styles = {
        "h1": ParagraphStyle(
            "ChineseHeading1",
            parent=sample["Heading1"],
            fontName="STSong-Light",
            fontSize=18,
            leading=24,
            spaceAfter=12,
        ),
        "h2": ParagraphStyle(
            "ChineseHeading2",
            parent=sample["Heading2"],
            fontName="STSong-Light",
            fontSize=14,
            leading=20,
            spaceAfter=8,
        ),
        "body": ParagraphStyle(
            "ChineseBody",
            parent=sample["BodyText"],
            fontName="STSong-Light",
            fontSize=10.5,
            leading=17,
            spaceAfter=8,
        ),
    }

    story = []
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            story.append(Spacer(1, 4))
            continue
        if line.startswith("# "):
            story.append(Paragraph(_escape(line[2:].strip()), styles["h1"]))
        elif line.startswith("## "):
            story.append(Paragraph(_escape(line[3:].strip()), styles["h2"]))
        elif line.startswith("- "):
            story.append(Paragraph(f"- {_escape(line[2:].strip())}", styles["body"]))
        elif line == "---":
            story.append(Spacer(1, 8))
        else:
            story.append(Paragraph(_escape(_strip_markdown_emphasis(line)), styles["body"]))

    document.build(story)
    return output.getvalue()


def _strip_markdown_emphasis(text: str) -> str:
    return text.replace("**", "").replace("`", "")


def _escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("**", "")
        .replace("`", "")
    )
